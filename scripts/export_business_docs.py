"""
Convert the two business markdown files to PDF + DOCX for client / partner.

PDF path: markdown → HTML (with print stylesheet) → headless Chrome → PDF.
DOCX path: walk the markdown AST with python-docx, render headings, paragraphs,
tables, images with reasonable defaults.

Output: docs/business/exports/

Run from the repo root:
    python scripts/export_business_docs.py
"""
from __future__ import annotations

import asyncio
import re
import sys
import tempfile
from pathlib import Path

import markdown as md_lib
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs" / "business"
EXPORT = DOCS / "exports"
EXPORT.mkdir(exist_ok=True)


# --- Shared HTML wrapper for the PDF route -----------------------------------

HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  @page {{ size: A4; margin: 18mm 16mm; }}
  body {{
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    color: #1a1a1a;
    font-size: 11pt;
    line-height: 1.55;
    max-width: 100%;
  }}
  h1 {{
    color: #E50914;
    font-size: 24pt;
    border-bottom: 3px solid #E50914;
    padding-bottom: 6px;
    margin-top: 32px;
    page-break-before: auto;
  }}
  h1:first-of-type {{ page-break-before: avoid; }}
  h2 {{
    color: #141414;
    font-size: 17pt;
    margin-top: 28px;
    border-bottom: 1px solid #ddd;
    padding-bottom: 4px;
  }}
  h3 {{
    color: #141414;
    font-size: 13pt;
    margin-top: 20px;
  }}
  h4 {{ color: #333; font-size: 11.5pt; margin-top: 14px; }}
  p {{ margin: 8px 0; }}
  table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 10pt;
  }}
  th {{
    background: #141414;
    color: white;
    text-align: left;
    padding: 6px 8px;
    font-weight: 600;
  }}
  td {{
    border: 1px solid #e5e5e5;
    padding: 6px 8px;
    vertical-align: top;
  }}
  tr:nth-child(even) td {{ background: #fafafa; }}
  code, pre {{
    font-family: "Consolas", "Menlo", monospace;
    font-size: 9.5pt;
    background: #f5f5f5;
    padding: 1px 4px;
    border-radius: 3px;
  }}
  pre {{ padding: 10px; overflow-x: auto; }}
  img {{
    max-width: 100%;
    height: auto;
    margin: 10px 0;
    border: 1px solid #e0e0e0;
    border-radius: 6px;
    page-break-inside: avoid;
  }}
  blockquote {{
    border-left: 3px solid #E50914;
    margin: 12px 0;
    padding: 4px 12px;
    color: #555;
    background: #fafafa;
  }}
  hr {{ border: 0; border-top: 1px solid #ddd; margin: 24px 0; }}
  ul, ol {{ margin: 6px 0 6px 22px; padding: 0; }}
  li {{ margin: 3px 0; }}
  /* Avoid awkward page breaks inside small blocks */
  table, h2, h3 {{ page-break-inside: avoid; }}
  /* Avoid orphan/widow headings */
  h2 + p, h3 + p {{ page-break-before: avoid; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""


def strip_frontmatter(text: str) -> tuple[str, dict]:
    """Pull off a YAML frontmatter block if present. Returns (body, meta)."""
    if not text.startswith("---"):
        return text, {}
    end = text.find("\n---", 4)
    if end < 0:
        return text, {}
    header = text[4:end].strip()
    body = text[end + 4 :].lstrip("\n")
    meta = {}
    for line in header.splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            meta[k.strip()] = v.strip().strip('"')
    return body, meta


def md_to_html(md_text: str) -> str:
    """Markdown → HTML using python-markdown's table + fenced-code extensions."""
    return md_lib.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists", "attr_list"],
    )


async def html_to_pdf(html: str, out_path: Path, base_dir: Path) -> None:
    """Render HTML with Playwright's headless Chromium and print to PDF.

    base_dir matters because the markdown references images via
    `screenshots/...` (relative). We write the HTML to a temp file inside
    base_dir so the browser resolves those relative URLs correctly.
    """
    from playwright.async_api import async_playwright

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", encoding="utf-8", dir=base_dir, delete=False
    ) as fh:
        fh.write(html)
        tmp = Path(fh.name)
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.goto(f"file://{tmp.as_posix()}", wait_until="networkidle")
            await page.pdf(
                path=str(out_path),
                format="A4",
                margin={"top": "18mm", "right": "16mm", "bottom": "18mm", "left": "16mm"},
                print_background=True,
            )
            await browser.close()
    finally:
        tmp.unlink(missing_ok=True)


# --- DOCX render -------------------------------------------------------------


def _set_cell_bg(cell, color_hex: str) -> None:
    """Apply a background fill to a DOCX table cell (no shortcut in python-docx)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def md_to_docx(md_text: str, out_path: Path, base_dir: Path) -> None:
    """Walk the markdown line-by-line (good enough for our docs) and emit DOCX."""
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Horizontal rule
        if re.fullmatch(r"-{3,}", line):
            p = doc.add_paragraph()
            p.add_run("_" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))  # strip **bold**
            h = doc.add_heading(text, level=min(level, 4))
            if level == 1:
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0xE5, 0x09, 0x14)
            i += 1
            continue

        # Image — markdown ![alt](path)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if m:
            img_path = (base_dir / m.group(2)).resolve()
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.0))
                except Exception:
                    doc.add_paragraph(f"[image: {m.group(2)} — failed to embed]")
            else:
                doc.add_paragraph(f"[image missing: {m.group(2)}]")
            i += 1
            continue

        # Table — detect `| ... |` line followed by `| --- | --- |` separator
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            i += 2  # skip separator
            data_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                # Pad / trim to match header width
                row = (row + [""] * len(header_cells))[: len(header_cells)]
                data_rows.append(row)
                i += 1

            table = doc.add_table(rows=1, cols=len(header_cells))
            table.style = "Light Grid"
            hdr = table.rows[0].cells
            for j, cell_text in enumerate(header_cells):
                hdr[j].text = cell_text
                _set_cell_bg(hdr[j], "141414")
                for p in hdr[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for r_idx, row in enumerate(data_rows):
                cells = table.add_row().cells
                for j, cell_text in enumerate(row):
                    cells[j].text = _strip_inline_md(cell_text)
                    if r_idx % 2 == 1:
                        _set_cell_bg(cells[j], "FAFAFA")
            doc.add_paragraph()  # spacer
            continue

        # Bullets — `- ` or `* `
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            doc.add_paragraph(_strip_inline_md(text), style="List Bullet")
            i += 1
            continue

        # Numbered list — `1. ` etc.
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            doc.add_paragraph(_strip_inline_md(text), style="List Number")
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — accumulate until blank line
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i]):
            para_lines.append(lines[i])
            i += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, text)

    doc.save(str(out_path))


def _is_block_start(line: str) -> bool:
    return (
        line.startswith("#")
        or line.startswith("|")
        or line.startswith("- ")
        or line.startswith("* ")
        or bool(re.match(r"^\d+\.\s", line))
        or line.startswith("![")
        or re.fullmatch(r"-{3,}", line.strip()) is not None
    )


def _strip_inline_md(text: str) -> str:
    """Drop *italic*, **bold**, `code`, and [link](url) wrappers — DOCX tables
    don't need rich runs, plain text is fine."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _add_inline_runs(para, text: str) -> None:
    """Split text on **bold**, *italic*, `code` and emit styled runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = para.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*"):
            r = para.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`"):
            r = para.add_run(token[1:-1])
            r.font.name = "Consolas"
        elif token.startswith("["):
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if inner:
                para.add_run(inner.group(1))
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


# --- Driver -------------------------------------------------------------------


async def main() -> None:
    sources = [
        ("partner-full-features.md", "CinePile — Full Feature Inventory"),
        ("client-pitch.md", "CinePile — Product Pitch"),
    ]
    for filename, _title in sources:
        src = DOCS / filename
        if not src.exists():
            print(f"  ! missing {src}; skip")
            continue

        text = src.read_text(encoding="utf-8")
        body, meta = strip_frontmatter(text)
        title = meta.get("title", filename)

        # PDF
        html = HTML_TEMPLATE.format(title=title, body=md_to_html(body))
        pdf_out = EXPORT / src.with_suffix(".pdf").name
        await html_to_pdf(html, pdf_out, base_dir=DOCS)
        size_kb = pdf_out.stat().st_size // 1024
        print(f"  PDF  -> {pdf_out.relative_to(ROOT)}  ({size_kb} KB)")

        # DOCX
        docx_out = EXPORT / src.with_suffix(".docx").name
        md_to_docx(body, docx_out, base_dir=DOCS)
        size_kb = docx_out.stat().st_size // 1024
        print(f"  DOCX -> {docx_out.relative_to(ROOT)}  ({size_kb} KB)")


if __name__ == "__main__":
    print("Exporting business docs to PDF + DOCX...")
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        sys.exit(130)
